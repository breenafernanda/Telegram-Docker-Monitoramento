
from PIL import Image, ImageTk
import docx
from docx.shared import Inches
import os
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from docx.shared import Cm
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
from docx.shared import Inches
import docx2pdf
from docx2pdf import convert
import time
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
import locale
import numpy as np
from datetime import datetime
import datetime
import shutil
from unidecode import unidecode

def gerar_pdf(file_path,palavras_especialista_excelente,palavras_especialista_abaixo):
            now = datetime.datetime.now()
            
            mes_ano = now.strftime('%m%Y')
            
            # Criar pasta PDFs se não existir
            pdfs_directory = criar_pasta_pdfs()
            print('=============PASSOU===============')

            # Criar pasta mes_ano se não existir
            mes_ano_directory = criar_pasta_mes_ano(pdfs_directory)
                        
            
            adicionar_primeira_linha(file_path)
            
            def gerar_grafico(datas, gerado, expectativa):
                try:
                    datas_Mes_Ano = []
                    # print(datas)
                    print(f'Intervalo de datas: {datas[0]} - {datas[-1]}\n\n')

                    for dia in datas:
                        dia_mes_ano = dia.split('/')
                        data_formatada = f"{dia_mes_ano[0]}/{dia_mes_ano[1]}"
                        datas_Mes_Ano.append(data_formatada)
                    # print(datas_Mes_Ano,'\n', gerado)
                    gerado_new = []
                    total_gerado = 0
                    dias = 0
                    for gerado_dia in gerado:
                        try:
                            gerado_dia = float(gerado_dia)
                        except:
                            gerado_dia = 0
                        gerado_new.append(float(gerado_dia))
                        total_gerado = total_gerado + gerado_dia
                        dias = dias + 1
                    
                    expectativa = (total_gerado * 100 / desempenho) * 0.85
                    linha_prognostica  = expectativa / dias
                    # print(f'Expec: `{expectativa:.1f}    ||    dias: {dias}    || progn: {prog:.1f}')
                    # Adicionando a linha de prognóstico (média)
                    # plt.axhline(prog, color="green", linestyle='--', label='Linha Prognóstica')
                    
                    plt.axhline(linha_prognostica , color=(0/255, 178/255, 188/255), linewidth=4,linestyle='-', label='Linha Prognóstica')
                    

                    # Selecionando apenas alguns índices para exibir no eixo de DATAS
                    indices_xticks = np.arange(0, len(datas_Mes_Ano), step=5)
                    # Configurando os xticks selecionados
                    plt.xticks(indices_xticks, [datas_Mes_Ano[i] for i in indices_xticks], rotation=90)

                        # plt.xticks(indices_xticks, [datas_Mes_Ano[i] for i in indices_xticks], rotation=90)
                    plt.rcParams['figure.facecolor'] = 'none'  


                    print(f'\x1b[34m\n\n----------------------------\nGerando gráfico: gerados: \x1b[33m{gerado_new}\n\x1b[34m----------------------------\n\n\x1b[0m')
                    plt.bar(datas_Mes_Ano, gerado_new, color='#BF4F8E')
                    plt.xlabel('Gerado(kWh)')
                    plt.ylabel('kWh/mês gerados')

                    # # Definir a cor de fundo do gráfico
                    plt.gca().set_facecolor('none') 
                    plt.gcf().set_facecolor('none')

                    # Adicionar título ao gráfico
                    plt.title('Gerado(kWh)')

                    # plt.show() #visualizar o grafico

                    plt.savefig('grafico_gerado.png', transparent=True)
                    plt.close()  # Fechar a figura para liberar memória
                    # Redimensionar a imagem
                    largura_desejada = 500  # em polegadas
                    altura_desejada = 280 # em polegadas
                    try:
                        img = Image.open('grafico_gerado.png')
                        img = img.resize((int(largura_desejada), int(altura_desejada)))  # Convertendo polegadas para pixels

                        img.save('grafico_gerado.png')
                    except Exception as e: print(f'Erro ao salvar o gráfico: {e}')
                except:
                    pass

            def gerar_relatorio(id, cliente, vetor_gerado, desempenho, mes,palavras_especialista_abaixo,palavras_especialista_excelente):
                try:
                        # print(f'{id}   -  {cliente}', end="   |  ")
                        total_gerado = 0
                        for gerado_dia in vetor_gerado:
                            if gerado_dia == 'None':
                                gerado_dia = 0
                            total_gerado = float(total_gerado) + float(gerado_dia)
       

                        print(f"total gerado: {total_gerado:.1f} kWh   \x1b[33m desempenho: {desempenho} %\x1b[0m", end=" ")
                        cem_por_cento = ( 100 * float(total_gerado) /float(desempenho) ) * 0.85

                        print(f'\x1b[32mExpectativa de geração: {cem_por_cento:.1f} kWh\x1b[36m', end=" ")
                        # novo desempenho
                        new_desempenho = (total_gerado*100)/cem_por_cento
                        
                        print(f'Novo desempenho = {new_desempenho}\x1b[0m') ## exemplo 105,25%% 
                        if new_desempenho >= 90:
                            palavras_do_especialista = palavras_especialista_excelente
                        elif new_desempenho >=72 and new_desempenho <90:
                            palavras_do_especialista = palavras_especialista_abaixo

                        print(f'PALAVRAS DO ESPECIALISTA ESCOLHIDA ESCOLHIDA:\x1b[33m{palavras_do_especialista}\x1b[0m')
                        
                        def atribuir_estilo_celula(cell, texto):
                            cor = RGBColor(92, 92, 92)  # Cor cinza chumbo (RGB: 32, 32, 32)
                            fonte = "Abadi"
                            tamanho = 18
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run(texto)
                            run.font.name = fonte
                            run.font.size = Pt(tamanho)
                            run.font.color.rgb = cor
                            run.font.bold = True  # Adiciona negrito (bold)
                        def atribuir_estilo_celula_2(cell, texto):
                            cor = RGBColor(255, 255, 255)  # Cor cinza chumbo (RGB: 32, 32, 32)
                            fonte = "Abadi"
                            tamanho = 30
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run(texto)
                            run.font.name = fonte
                            run.font.size = Pt(tamanho)
                            run.font.color.rgb = cor
                            run.font.bold = True  # Adiciona negrito (bold)
                        def atribuir_estilo_celula_2RS(cell, texto):
                            cor = RGBColor(255, 255, 255)  # Cor cinza chumbo (RGB: 32, 32, 32)
                            fonte = "Abadi"
                            tamanho = 28
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run(texto)
                            run.font.name = fonte
                            run.font.size = Pt(tamanho)
                            run.font.color.rgb = cor
                            run.font.bold = True  # Adiciona negrito (bold) 
                        def atribuir_estilo_celula_3(cell, texto):
                            cor = RGBColor(255, 255, 255)  # Cor cinza chumbo (RGB: 32, 32, 32)
                            fonte = "Abadi"
                            tamanho = 14
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run(texto)
                            run.font.name = fonte
                            run.font.size = Pt(tamanho)
                            run.font.color.rgb = cor
                            run.font.bold = True  # Adiciona negrito (bold)
                        def atribuir_estilo_celula_4(cell, texto):
                            cor = RGBColor(176, 30, 126)  # Cor cinza chumbo (RGB: 32, 32, 32)
                            fonte = "Abadi"
                            tamanho = 20
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run(texto)
                            run.font.name = fonte
                            run.font.size = Pt(tamanho)
                            run.font.color.rgb = cor
                            run.font.bold = True  # Adiciona negrito (bold)
                        def atribuir_estilo_celula_5(cell, texto):
                            cor = RGBColor(255, 255, 255)  # Cor cinza chumbo (RGB: 32, 32, 32)
                            fonte = "Abadi"
                            tamanho = 22
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run(texto)
                            run.font.name = fonte
                            run.font.size = Pt(tamanho)
                            run.font.color.rgb = cor
                            run.font.bold = False  # Adiciona negrito (bold)
                        def atribuir_estilo_celula_6(cell, texto):
                            cor = RGBColor(255, 255, 255)  # Cor cinza chumbo (RGB: 32, 32, 32)
                            fonte = "Abadi"
                            tamanho = 22
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run(texto)
                            run.font.name = fonte
                            run.font.size = Pt(tamanho)
                            run.font.color.rgb = cor
                            run.font.bold = True  # Adiciona negrito (bold)
                        
                        def salvar():
                            try:
                                caminho_salvar = f'Relatorio de Geração.docx'
                                # caminho_salvar = os.path.join('configs',caminho_salvar)

                                document.save(caminho_salvar)
                                # imprimir(f'\nDocumento salvo como: ', "black")
                                # imprimir(f'Analise de Fatura de Energia.docx', "black")

                            except: 
                                # imprimir(f'Erro ao salvar, fechar arquivo e tentar novamente')
                                try:
                                    i = 1
                                    aberto = True
                                    while (aberto==True):
                                        try:
                                            caminho_salvar = f'Relatorio de Geração({i}).docx'
                                            # caminho_salvar = os.path.join('configs',caminho_salvar)
                                            document.save(caminho_salvar)
                                            aberto = False
                                            # imprimir(f'\nDocumento salvo como: ', "black")
                                            # imprimir(f'Analise de Fatura de Energia ({i}).docx', "black")
                                            aberto = '✘'
                            #                 i=0
                                        except:
                                            i = i + 1
                                except Exception as e: print(e)
                            return caminho_salvar
                        
                        # carregar o arquivo de modelo de relatório e abrir o document
                        arquivo_modelo_relatorio = os.path.join('configs',"Modelo_relatorio_geracao.docx")
                        document = docx.Document(arquivo_modelo_relatorio)

                        # definir celulas das tabelas para inserir as informações
                        cell_vigencia = document.tables[2].cell(0, 0)
                        
                        cell_grafico = document.tables[9].cell(0, 0)  # Acessa a célula desejada (linha 0, coluna 0)
                        cell_mes_voce_gerou = document.tables[6].cell(0, 0)
                        cell_mes_nos_esperavamos = document.tables[7].cell(0, 0)
                        cell_mes_representa_economia= document.tables[8].cell(0, 0)
                        cell_kg_co2= document.tables[10].cell(0, 0)
                        cell_arvores= document.tables[11].cell(0, 0)
                        cell_km= document.tables[12].cell(0, 0)
                        cell_cliente= document.tables[13].cell(0, 0)
                        cell_especialista= document.tables[14].cell(0, 0)

                        economia_calculada = total_gerado * 1.12
                        co2 = total_gerado * 0.475
                        arvores = total_gerado * 0.025956
                        km = total_gerado * 7.5

                        representa_economia = document.tables[4].cell(0, 0)
                        voce_gerou = document.tables[3].cell(0, 0)
                        nos_esperavamos = document.tables[5].cell(0, 0)
                        # print(total_gerado, 'total gerado')
                        # print(economia_calculada, 'economia')
                        # print(cem_por_cento, 'cemPorCento')
                        # inserir conteudo nas celulas de geração
                        atribuir_estilo_celula(cell_vigencia, f'{mes}') # campo de mês vigente
                        atribuir_estilo_celula_2(voce_gerou, f'{total_gerado:.1f}') # campo de "Você gerou xxx kWH ''
                        atribuir_estilo_celula_2(nos_esperavamos, f'{cem_por_cento:.1f}') # campo de "Nós esperavamos xxx kWh "
                        atribuir_estilo_celula_2RS(representa_economia, f'{economia_calculada:.2f}') # verificar equação para estimar economia
                        atribuir_estilo_celula_3(cell_mes_voce_gerou, f'{mes}')
                        atribuir_estilo_celula_3(cell_mes_nos_esperavamos,  f'{mes}')
                        atribuir_estilo_celula_3(cell_mes_representa_economia,  f'{mes}')
                        atribuir_estilo_celula_4(cell_kg_co2,  f'{co2:.0f}')
                        atribuir_estilo_celula_4(cell_arvores,  f'{arvores:.0f}')
                        atribuir_estilo_celula_4(cell_km,  f'{km:.0f}')
                        atribuir_estilo_celula_5(cell_cliente,  f'{cliente}\n')
                        atribuir_estilo_celula_6(cell_especialista,  f'{palavras_do_especialista}')

                        # Localize a célula para inserir o gráfico
                        cell_grafico_paragraph = cell_grafico.paragraphs[0]  # Obtém o parágrafo da célula
                        run = cell_grafico_paragraph.add_run()
                        run.add_picture("grafico_gerado.png")  # Insere a imagem na célula

                        doc = salvar()
                        # # Load word document
                        ## tirar acentos do nome 
                        try:


                            def remover_acentos(texto):
                                try:
                                    texto_sem_acentos = unidecode(texto)
                                    return texto_sem_acentos
                                except Exception as e:
                                    raise ValueError("Erro ao remover acentos: {}".format(e))
                            cliente = remover_acentos(cliente)
                        except Exception as e: print(f'>>< {e}')

                        doc_pdf = f"{cliente}.pdf"
                        
                        
                        # Converter e salvar como PDF
                        convert(doc, doc_pdf)

                        # Obtém o diretório atual
                        path_atual = os.getcwd()

                        # Criar pasta mes_ano se não existir
                        # mes_ano_directory = criar_pasta_mes_ano(pdfs_directory, mes_ano)
                        mes_ano_directory = os.path.join(pdfs_directory, mes_ano)
                        print('mes_ano_directory: ', mes_ano_directory)
                        ## mover arquivo do diretorio atual para a pasta com mes e ano correspondente dentro da pasta PDFs
                        ## doc_pdf to mes_ano_directory
                        caminho_origem = os.path.join(path_atual, doc_pdf)

                        # caminho_destino = mes_ano_directory + 'doc_pdf'
                        caminho_destino = os.path.join(mes_ano_directory, doc_pdf)

                        # Mover o arquivo da origem para o destino
                        shutil.move(caminho_origem, caminho_destino)

                        print(f"Arquivo movido para: {caminho_destino}")
                        # Imprimir mensagem de confirmação
                        print(f'PDF gerado e salvo em: {mes_ano_directory}\{doc_pdf}')
                        
                except:
                    pass
            
            def zipar_relatorio(pdfs_directory, mes_ano):
                print('pasta pdf: ',pdfs_directory)
                mes_ano_directory = os.path.join(pdfs_directory, mes_ano)

                # Constrói o caminho para o arquivo ZIP
                zip_filename = f"Relatórios de Geração - {mes_ano}.zip"  # Nome do arquivo ZIP
                caminho_zip = os.path.join(mes_ano_directory, zip_filename)
                # print(mes_ano_directory, '\n',caminho_zip)
                listar_arquivos = os.listdir(mes_ano_directory)
                # for arquivo in listar_arquivos:
                #     print('arquivo: ', arquivo)
                # Lista todos os arquivos na pasta mes_ano_directory
                arquivos = listar_arquivos
                print(arquivos)
                # Cria o arquivo ZIP e adiciona os arquivos
                try:
                    with zipfile.ZipFile(caminho_zip, 'w') as zipf:
                        for arquivo in arquivos:
                            try:
                                print(arquivo)
                                arquivo = os.path.join(mes_ano_directory, arquivo)
                                print(arquivo)

                                zipf.write(arquivo, os.path.basename(arquivo))
                            except Exception as e: print('>>>',e)
                except Exception as e: print(e)
                print(f"\n\n---------------\nArquivo ZIP criado em: {caminho_zip}")
                return caminho_zip
            
            #  Ler o arquivo XLSX usando o pandas
            df = pd.read_excel(file_path)
            # Obter o número de colunas
            num_colunas = len(df.columns)
            # Criar um vetor com todas as colunas
            colunas = df.columns.tolist()
            # Imprimir  número de colunas
            print(f"O arquivo tem {num_colunas} usinas.    ")
            # Imprimir os valores específicos de cada coluna
            vetor_dia = []
            vetor_gerado = []
            dados_geracao = []

            usinas_90 = []
            usinas_70 = []
            usinas_nao = []
            id = 0

            for coluna in colunas:
                id += 1
                try:
                    dados_geracao.clear()
                    vetor_dia.clear()
                    vetor_gerado.clear()
                    # Obter os valores da coluna como uma lista
                    valores_coluna = df[coluna].values.tolist()
                    # Salvar os valores da linha 4 até a última linha
                    dados_geracao += valores_coluna[6:]
                    # print(dados_geracao)
                    cliente = valores_coluna[0]
                    email = valores_coluna[1]
                    data_instalacao = valores_coluna[2]
                    desempenho = valores_coluna[3]
                    mes = valores_coluna[4]
                    print(f"\n-------\nCliente: {cliente}\nE-mail: {email}\nData de Instalação: {data_instalacao}\nDesempenho: {desempenho}\nMês de Referência: {mes}\n")  # Print the contents of each linha
                    try:
                        desempenho = desempenho.replace('%', '').replace(',','.')
                    except:pass
                    try:
                        desempenho = float(desempenho)
                    except: pass
                    usina_dados = [cliente, desempenho, email]
                    for dado in dados_geracao:
                        try:
                            valores = dado.split(",")
                            vetor_dia.append(valores[0])

                            if valores[1] != 'nan' or valores[1] != None:
                                vetor_gerado.append(float(valores[1]))
                            else:
                                vetor_gerado.append(0)
                                # desempenho = 0
                        except:
                            vetor_dia.append('None')
                            vetor_gerado.append('None')
                
                    # print(vetor_gerado, desempenho)
                
                    if desempenho >= 90:
                        palavras_do_especialista = palavras_especialista_excelente
                        print(f'\x1b[36m\n>>>        {desempenho} %   {valores_coluna[0]}\x1b[0m')
                        gerar = True
                        usinas_90.append(usina_dados)
                    elif desempenho >=72 and desempenho <90:
                        palavras_do_especialista = palavras_especialista_abaixo
                        print(f'\x1b[33m\n>>>        {desempenho} %   {valores_coluna[0]}\x1b[0m')
                        gerar = True
                        usinas_70.append(usina_dados)
                    else: 
                        print(f'{desempenho:.0f} %   {valores_coluna[0]}')
                        print(f'\x1b[32m\n>>> {id}    -  RELATORIO NÃO GERADO: desempenho: {desempenho:.0f} %   ||   {cliente} \x1b[0m')
                        gerar = False
                        usinas_nao.append(usina_dados)

                    if gerar == True:
                        try:
                            # print(vetor_gerado)
                            try:
                                gerar_grafico(vetor_dia, vetor_gerado, cliente)
                            except Exception as e: print(f'Erro ao gerar gráfico ** : {e}')
                            gerar_relatorio(id,cliente, vetor_gerado, desempenho, mes, palavras_especialista_abaixo,palavras_especialista_excelente)
                            print('\n---------------------------------------\n')
                        except Exception as e: print(e)


                # desempenho = f'{desempenho:.2f}'
                except Exception as e: print(e)


            # caminho_zip_file = zipar_relatorio(pdfs_directory, mes_ano)    
            # print('\n\ncaminho_zip_file: ', caminho_zip_file)
            remover_primeira_linha(file_path)



def criar_pasta_pdfs():
    current_directory = os.path.dirname(os.path.abspath(__file__))

    pdfs_directory = os.path.join(current_directory,'PDFs')
    if not os.path.exists(pdfs_directory):
        os.mkdir(pdfs_directory)
    return pdfs_directory

def criar_pasta_mes_ano(pdfs_directory):
    now = datetime.datetime.now()
    mes_ano = now.strftime('%m%Y')
    mes_ano_directory = os.path.join(pdfs_directory, mes_ano)
    if not os.path.exists(mes_ano_directory):
        os.mkdir(mes_ano_directory)
    return mes_ano_directory

# por algum motivo só estava lendo da linha 2 pra baixo
#  então ao inicio da execução adiciona uma linha no topo e no final a remove 
def adicionar_primeira_linha(file_path):
    # Carregar o arquivo XLSX
    workbook = openpyxl.load_workbook(file_path)

    # Selecionar a primeira planilha
    sheet = workbook.active

    # Inserir uma linha em branco no topo da planilha
    sheet.insert_rows(1)

    # Salvar as alterações no arquivo
    workbook.save(file_path)

def remover_primeira_linha(file_path):
    # Carregar o arquivo XLSX
    workbook = openpyxl.load_workbook(file_path)

    # Selecionar a primeira planilha
    sheet = workbook.active

    # Remover a primeira linha da planilha
    sheet.delete_rows(1)

    # Salvar as alterações no arquivo
    workbook.save(file_path)

